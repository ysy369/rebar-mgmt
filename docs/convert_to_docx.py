# 将操作手册 Markdown 转换为 WPS 兼容的 DOCX 格式
import re, os

from docx import Document
from docx.shared import Pt, RGBColor

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Microsoft YaHei"
style.font.size = Pt(11)

src = os.path.join(os.path.dirname(__file__), "07-操作手册.md")
with open(src, "r", encoding="utf-8") as f:
    lines = f.readlines()

for line in lines:
    orig = line.rstrip()
    if not orig:
        continue

    if orig.startswith("# ") and not orig.startswith("## "):
        doc.add_heading(orig[2:], level=1)
    elif orig.startswith("## "):
        doc.add_heading(orig[3:], level=2)
    elif orig.startswith("### "):
        doc.add_heading(orig[4:], level=3)
    elif orig.startswith("|"):
        continue  # skip table (too complex for quick conversion)
    elif orig.startswith("- "):
        doc.add_paragraph(orig[2:], style="List Bullet")
    elif re.match(r"^> ", orig):
        p = doc.add_paragraph()
        run = p.add_run(orig[2:])
        run.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)
    elif orig.startswith("```"):
        continue
    else:
        # inline bold
        p = doc.add_paragraph()
        parts = re.split(r"(\*\*.*?\*\*|`.*?`)", orig)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                r = p.add_run(part[2:-2])
                r.bold = True
            elif part.startswith("`") and part.endswith("`"):
                r = p.add_run(part[1:-1])
                r.font.name = "Consolas"
            else:
                p.add_run(part)

dst = os.path.join(os.path.dirname(__file__), "钢筋精细化管理平台-操作手册.docx")
doc.save(dst)
print(f"DOCX saved: {dst}")
